import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn
import datetime

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=80, bottom=80, left=120, right=120):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(f'''
        <w:tcMar {nsdecls("w")}>
            <w:top w:w="{top}" w:type="dxa"/>
            <w:bottom w:w="{bottom}" w:type="dxa"/>
            <w:left w:w="{left}" w:type="dxa"/>
            <w:right w:w="{right}" w:type="dxa"/>
        </w:tcMar>
    ''')
    tcPr.append(tcMar)

def set_table_borders(table):
    tblPr = table._tbl.tblPr
    borders = parse_xml(f'''
        <w:tblBorders {nsdecls("w")}>
            <w:top w:val="single" w:sz="6" w:space="0" w:color="D3D3D3"/>
            <w:left w:val="none"/>
            <w:bottom w:val="single" w:sz="8" w:space="0" w:color="1B365D"/>
            <w:right w:val="none"/>
            <w:insideH w:val="single" w:sz="4" w:space="0" w:color="E0E0E0"/>
            <w:insideV w:val="none"/>
        </w:tblBorders>
    ''')
    tblPr.append(borders)

def create_printable_document():
    doc = docx.Document()
    
    # 1. Setup Page Size to A4 & Margins (Optimized to fit strictly on 1 Page)
    section = doc.sections[0]
    section.page_width = Inches(8.27)   # A4 Width
    section.page_height = Inches(11.69) # A4 Height
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.6)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)

    # Styles
    styles = doc.styles
    normal_style = styles['Normal']
    normal_style.font.name = 'Calibri'
    normal_style.font.size = Pt(10)
    normal_style.font.color.rgb = RGBColor(0x22, 0x22, 0x22)

    # Header / Title Block
    title_p = doc.add_paragraph()
    title_p.paragraph_format.space_before = Pt(0)
    title_p.paragraph_format.space_after = Pt(2)
    run_title = title_p.add_run("ORÇAMENTO DE PEÇAS - MANUTENÇÃO DE NOTEBOOK")
    run_title.font.size = Pt(16)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(0x1B, 0x36, 0x5D)

    # Client Info & Date Block Table (Print Friendly)
    info_table = doc.add_table(rows=2, cols=2)
    info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    info_table.autofit = False
    
    info_table.rows[0].cells[0].width = Inches(4.5)
    info_table.rows[0].cells[1].width = Inches(2.4)
    info_table.rows[1].cells[0].width = Inches(4.5)
    info_table.rows[1].cells[1].width = Inches(2.4)

    c0 = info_table.rows[0].cells[0].paragraphs[0].add_run("Cliente: _____________________________________")
    c0.font.size = Pt(9.5)
    c0.font.bold = True
    
    c1 = info_table.rows[0].cells[1].paragraphs[0].add_run(f"Data: {datetime.datetime.now().strftime('%d/%m/%Y')}")
    c1.font.size = Pt(9.5)
    info_table.rows[0].cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

    c2 = info_table.rows[1].cells[0].paragraphs[0].add_run("Equipamento: Notebook CCE / Dell (Manutenção)")
    c2.font.size = Pt(9.5)

    c3 = info_table.rows[1].cells[1].paragraphs[0].add_run("Validade da Proposta: 5 dias")
    c3.font.size = Pt(9.5)
    info_table.rows[1].cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Divider line
    p_div = doc.add_paragraph()
    p_div.paragraph_format.space_before = Pt(4)
    p_div.paragraph_format.space_after = Pt(8)
    p_div_border = parse_xml(f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="12" w:space="1" w:color="1B365D"/></w:pBdr>')
    p_div._p.get_or_add_pPr().append(p_div_border)

    # Section Heading
    sec_p = doc.add_paragraph()
    sec_p.paragraph_format.space_after = Pt(6)
    r_sec = sec_p.add_run("1. Relação de Componentes e Peças")
    r_sec.font.bold = True
    r_sec.font.size = Pt(11)
    r_sec.font.color.rgb = RGBColor(0x1B, 0x36, 0x5D)

    # Table of Items
    table = doc.add_table(rows=1, cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table)

    col_widths = [Inches(0.6), Inches(3.6), Inches(1.1), Inches(0.4), Inches(1.2)]

    # Header row
    hdr_cells = table.rows[0].cells
    headers = ["Item", "Descrição da Peça", "Status", "Qtd", "Valor Unit."]
    for i, title in enumerate(headers):
        hdr_cells[i].text = title
        set_cell_background(hdr_cells[i], "1B365D")
        set_cell_margins(hdr_cells[i], top=100, bottom=100, left=80, right=80)
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i in [0, 2, 3, 4] else WD_ALIGN_PARAGRAPH.LEFT
        for run in p.runs:
            run.font.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            run.font.size = Pt(9.5)

    # Data Rows
    items = [
        ("01", "Tela 14.0\" LED Notebook CCE Ultra Thin U25 (LP140WH4 / B140XW01)", "Disponível", "1", "R$ 179,99", False),
        ("02", "Bateria Compatível Dell/CCE CL341-TS23 (SU341 / 3 Células 33Wh 7.4V)", "Disponível", "1", "R$ 342,62", False),
        ("03", "SSD Kingston A400 240GB 2.5\" SATA III", "Disponível", "1", "R$ 346,75", False),
        ("04", "Bateria CCE Original Ultra Thin T345 CL341-TS23", "Indisponível", "1", "—", True),
        ("05", "Memória RAM Smart DDR3L 8GB 1.35V PC3L-12800S", "Disponível", "1", "R$ 299,99", False),
    ]

    for idx, (item_num, desc, status, qtd, valor, is_indisponivel) in enumerate(items):
        row_cells = table.add_row().cells
        bg_color = "F8F9FA" if idx % 2 == 1 else "FFFFFF"
        
        row_data = [item_num, desc, status, qtd, valor]
        for i, val in enumerate(row_data):
            row_cells[i].text = val
            set_cell_background(row_cells[i], bg_color)
            set_cell_margins(row_cells[i], top=80, bottom=80, left=80, right=80)
            p = row_cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i in [0, 2, 3] else (WD_ALIGN_PARAGRAPH.RIGHT if i == 4 else WD_ALIGN_PARAGRAPH.LEFT)
            
            for run in p.runs:
                run.font.size = Pt(9)
                if is_indisponivel and i == 2:
                    run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00) # Red
                    run.font.bold = True
                elif not is_indisponivel and i == 2:
                    run.font.color.rgb = RGBColor(0x1E, 0x7E, 0x34) # Green

    # Apply widths to all cells
    for row in table.rows:
        for i, w in enumerate(col_widths):
            row.cells[i].width = w

    # Summary Box Table
    p_sum_space = doc.add_paragraph()
    p_sum_space.paragraph_format.space_before = Pt(6)
    p_sum_space.paragraph_format.space_after = Pt(2)

    summary_table = doc.add_table(rows=3, cols=2)
    summary_table.alignment = WD_TABLE_ALIGNMENT.RIGHT
    summary_table.autofit = False
    
    sum_widths = [Inches(3.2), Inches(1.8)]
    
    summary_data = [
        ("Subtotal (Itens Disponíveis):", "R$ 1.169,35", False),
        ("Frete / Taxa de Envio (Bateria Item 02):", "R$ 5,66", False),
        ("VALOR TOTAL DO ORÇAMENTO:", "R$ 1.175,01", True),
    ]

    for idx, (label, val, is_total) in enumerate(summary_data):
        row_cells = summary_table.rows[idx].cells
        row_cells[0].text = label
        row_cells[1].text = val
        
        bg = "D9E1F2" if is_total else "F2F2F2"
        for i in range(2):
            set_cell_background(row_cells[i], bg)
            set_cell_margins(row_cells[i], top=60, bottom=60, left=100, right=100)
            p = row_cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            for run in p.runs:
                run.font.size = Pt(10.5 if is_total else 9)
                run.font.bold = True if is_total else False
                if is_total:
                    run.font.color.rgb = RGBColor(0x1B, 0x36, 0x5D)

    for row in summary_table.rows:
        for i, w in enumerate(sum_widths):
            row.cells[i].width = w

    # Section 2: Observações para Impressão
    sec2_p = doc.add_paragraph()
    sec2_p.paragraph_format.space_before = Pt(10)
    sec2_p.paragraph_format.space_after = Pt(4)
    r_sec2 = sec2_p.add_run("2. Observações Importantes")
    r_sec2.font.bold = True
    r_sec2.font.size = Pt(11)
    r_sec2.font.color.rgb = RGBColor(0x1B, 0x36, 0x5D)

    notes = [
        "A bateria CCE original (Item 04) encontra-se indisponível no fornecedor. Foi inclusa no valor total a bateria compatível (Item 02).",
        "A bateria compatível (Item 02) possui prazo estimado de entrega de 15 a 20 dias úteis e envio internacional sujeito a eventuais taxas dos Correios/Alfândega.",
        "Orçamento válido por 5 dias corridos a contar da data de emissão."
    ]

    for note in notes:
        np = doc.add_paragraph(style='List Bullet')
        np.paragraph_format.space_before = Pt(0)
        np.paragraph_format.space_after = Pt(2)
        run_n = np.add_run(note)
        run_n.font.size = Pt(8.5)
        run_n.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    # Signature Block (Perfect for Print out & Approval)
    sig_p = doc.add_paragraph()
    sig_p.paragraph_format.space_before = Pt(28)
    sig_p.paragraph_format.space_after = Pt(0)
    sig_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    r_sig = sig_p.add_run("____________________________________________________\nAssinatura do Cliente / Aprovado por")
    r_sig.font.size = Pt(9)
    r_sig.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    # Save document
    output_path = r"c:\PROJETOS\loja-roupas\Orcamento_Pecas_Notebook_Impressao.docx"
    doc.save(output_path)
    print(f"Print-optimized document saved successfully at: {output_path}")

if __name__ == "__main__":
    create_printable_document()
